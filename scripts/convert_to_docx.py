"""Convert the OEA manuscript from LaTeX to Word (.docx) for AMLA 2026 submission."""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from pathlib import Path
import re

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "oea-submission.docx"

doc = Document()

# ── Page setup (letter, 1" margins) ──────────────────────────────────────────
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(10)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.15

# ── Helpers ──────────────────────────────────────────────────────────────────

def add_title(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.size = Pt(16)
    run.bold = True

def add_author(name, affiliation, email=""):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(name)
    run.font.size = Pt(12)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(affiliation)
    run2.font.size = Pt(10)
    run2.italic = True

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.name = "Times New Roman"

def add_body(text):
    p = doc.add_paragraph(text)
    p.paragraph_format.first_line_indent = Pt(18)
    return p

def add_body_no_indent(text):
    return doc.add_paragraph(text)

def add_bullet(text):
    doc.add_paragraph(text, style="List Bullet")

def add_equation(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(10)

def add_table(headers, rows, caption=""):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.name = "Times New Roman"
    # Data
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(9)
                    run.font.name = "Times New Roman"
    if caption:
        p = doc.add_paragraph()
        run = p.add_run(caption)
        run.font.size = Pt(9)
        run.italic = True
    return table

def add_figure(path, caption, width=Inches(5.5)):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if Path(path).exists():
        # Convert PDF to PNG would be needed; for now add placeholder
        run = p.add_run(f"[Figure: {Path(path).name}]")
        run.italic = True
    else:
        run = p.add_run(f"[Figure: {Path(path).name}]")
        run.italic = True
    if caption:
        cp = doc.add_paragraph()
        run = cp.add_run(caption)
        run.font.size = Pt(9)
        run.italic = True

# ══════════════════════════════════════════════════════════════════════════════
# BUILD THE DOCUMENT
# ══════════════════════════════════════════════════════════════════════════════

add_title("OEA: Structured Recursive Calibration\nfor Generative Stability")
add_author("Tristen Pierson", "BitConcepts Research")

# ── Abstract ─────────────────────────────────────────────────────────────────
add_heading("Abstract", level=1)
add_body_no_indent(
    'We present an empirical study of the OEA (Ontology, Epistemic, Agentic) framework as a '
    'three-layer protocol for improving reliability in recursive generative systems under '
    'synthetic-data exposure. The central thesis is: recursive generative stability depends '
    'more strongly on calibration direction and epistemic filtering than on unconstrained '
    'retrieval augmentation. "Ontology" in this work refers to structured distributional '
    'anchoring\u2014not philosophical ontology. "Agentic" refers to recursive persistence '
    'dynamics\u2014not autonomous agency. We report four experiments: (1) a bigram-proxy ablation '
    'study across 12 variants, (2) a four-model real LLM validation on distilgpt2 (82M), '
    'gpt2 (124M), gpt-neo-125M (125M), and Qwen2.5-1.5B (1.5B, modern 2024 architecture) '
    'with BM25 retrieval, (3) a 30-step recursive memory drift benchmark, and (4) a baseline '
    'competition against temperature reduction, top-k, entropy filtering, and repetition '
    'penalty. The full OEA protocol achieves TRR = 0.836 and FRR = 0.081 (d = 4.56, '
    'p < 0.001). The anti-calibrated variant provides strong ablation evidence: inverting the '
    'selection signal degrades log-probability by \u22120.55 to \u22121.37 nats, while full OEA '
    'improves it by +0.62 to +1.63 nats across all four models and three architecture families. '
    'Results are fully reproducible.'
)

# ── Keywords ─────────────────────────────────────────────────────────────────
p = doc.add_paragraph()
run = p.add_run("Keywords: ")
run.bold = True
run.font.size = Pt(10)
p.add_run("recursive calibration, epistemic filtering, model collapse, language models, "
          "distributional stability, retrieval augmented generation")

# ── 1. Introduction ──────────────────────────────────────────────────────────
add_heading("1. Introduction", level=1)
add_body(
    "This work argues that recursive generative stability depends more strongly on "
    "calibration direction and epistemic filtering than on unconstrained retrieval augmentation. "
    "Retrieval without directional calibration is associated with harm to distributional "
    "fidelity; miscalibration actively reverses discrimination performance."
)
add_body(
    "Model collapse has become a central concern in recursive synthetic training pipelines, "
    "with evidence of degradation under specific feedback-loop assumptions [1, 2]. Subsequent "
    "studies show collapse behavior depends strongly on data-mixing assumptions, curation "
    "policy, and accumulation dynamics [3, 4]. This motivates a controlled framework that "
    "separates what is known, what is hypothesized, and what is empirically testable."
)
add_body(
    "The OEA framework provides such a structure. Ontological Anchoring defines boundary "
    "conditions on the distributional vocabulary\u2014not formal symbolic ontology. Epistemic "
    "Filtering introduces explicit falsification and calibration checks for generated claims. "
    "Recursive Feedback operationalizes interventions in measurable recursive loops\u2014not "
    "autonomous agency. This aligns with a Popperian criterion: claims must be framed so "
    "they can fail under evidence [5]."
)

add_heading("1.1 Scope and Non-Claims", level=2)
add_body_no_indent("This work makes the following explicit non-claims:")
add_bullet("No claim of artificial general intelligence or any step toward it.")
add_bullet('No claim of formal symbolic ontology; "Ontology" is used in its computational sense.')
add_bullet('"Agentic" refers to recursive persistence dynamics, not autonomous agency.')
add_bullet("No claim of causal proof; results indicate association and ablation consistency.")
add_bullet("No claim of general intelligence improvement beyond the experimental regime.")
add_body(
    "This work does not claim to provide a theory of consciousness, autonomous cognition, "
    "symbolic ontology, or general intelligence. The contribution is limited to substrate-level "
    "mechanisms for recursive generative stability under synthetic exposure."
)

# ── 2. Core Hypotheses ───────────────────────────────────────────────────────
add_heading("2. Core Hypotheses and Falsifiability", level=1)
add_body_no_indent(
    "H1 (Calibration Direction): Correct-direction calibration improves recursive "
    "distributional fidelity relative to unconstrained sampling."
)
add_body_no_indent(
    "H2 (Miscalibration Reversal): Inverting the calibration signal reverses discrimination: "
    "TRR falls below control and FRR rises above control."
)
add_body_no_indent(
    "H3 (RAG without Filter): Retrieval augmentation without epistemic filtering is associated "
    "with degraded distributional fidelity relative to full OEA."
)
add_body_no_indent(
    "H4 (Memory Drift): OEA-controlled recursion is associated with lower semantic drift "
    "than unconstrained recursion over long chains."
)

add_heading("What would falsify OEA?", level=2)
add_bullet("H1: oea_anchored log-probability is not consistently higher than control across all four models.")
add_bullet("H2: oea_miscalibrated TRR is not lower than control TRR, or FRR is not higher.")
add_bullet("H3: oea_rag_only log-probability is not lower than oea_anchored.")
add_bullet("H4: OEA-controlled semantic drift (JSD) is not lower than uncontrolled at step 30.")

# ── 3. Notation ──────────────────────────────────────────────────────────────
add_heading("3. Notation and Formal Definitions", level=1)
add_table(
    ["Symbol", "Definition"],
    [
        ["x_t", "Token sequence at recursive step t"],
        ["P_\u03b8", "Generative model with frozen parameters \u03b8"],
        ["A(\u00b7)", "Ontological anchoring operator (vocabulary projection)"],
        ["E(\u00b7)", "Epistemic filter operator (candidate selection)"],
        ["G_0", "Reference frozen model (original, not updated)"],
        ["K", "Number of epistemic filter candidates (K=3)"],
        ["TRR", "True rejection rate (higher = better discrimination)"],
        ["FRR", "False rejection rate (lower = better specificity)"],
        ["JSD(p,q)", "Jensen-Shannon divergence"],
        ["R_t", "Retrieval context at step t (BM25 passage)"],
        ["S_t", "Recursive stability: 1 \u2212 JSD(p_{x_t}, p_{x_0})"],
        ["\u03c4", "Dynamic log-prob threshold"],
    ],
    "Table 1. Notation used throughout this paper."
)

add_equation("x_{t+1} ~ P_\u03b8(x | A(x_t), E(x_t), R_t)")
add_body_no_indent("where R_t is the retrieval context at step t (BM25 passage prepended to prompt).")
add_equation("S_t = 1 \u2212 JSD(p_{x_t}, p_{x_0})")
add_body_no_indent("Higher S_t indicates the generated distribution remains closer to the seed. S_0 = 1.")
add_equation("E(x) = argmax_{c \u2208 C_K} log G_0(c | x)")

# ── 4. Related Work ──────────────────────────────────────────────────────────
add_heading("4. Related Work", level=1)
add_heading("4.1 Recursive Synthetic Training and Collapse", level=2)
add_body(
    "Empirical and theoretical analyses report failure modes in recursive loops [1, 2]. "
    "Mitigation regimes exist when real data are retained [3]. Drayson et al. [6] prevent "
    "collapse at training time via machine-generated text detection. Zhu et al. [7] prevent "
    "collapse via token-level corpus editing. Kov\u00e1\u010d et al. [8] show lexical diversity "
    "amplifies distributional shift. Keisha et al. [9] characterize knowledge collapse as "
    "factual accuracy eroding while surface fluency persists."
)
add_heading("4.2 RAG, Calibration, and Trust-Aware Orchestration", level=2)
add_body(
    "RAG methods improve grounding [10] but retrieval alone does not eliminate faithfulness "
    "errors. Abbasi Yadkori et al. [11] provide an information-theoretic separation of "
    "epistemic from aleatoric uncertainty. Calibration diagnostics [12] are used as monitored "
    "variables, not truth proxies."
)
add_heading("4.3 Position Against Related Frameworks", level=2)
add_body(
    "OEA differs from constitutional AI, self-refinement, self-consistency, and recursive "
    "self-improvement by operating as a generation-time, frozen-weights, calibration-directed "
    "protocol. Unlike RAG-only methods, OEA studies the interaction between retrieval, "
    "directional calibration, and recursive filtering. Unlike generic decoding constraints, "
    "it evaluates whether directionally aligned calibration improves epistemic discrimination "
    "under recursive exposure."
)

# ── 5. Operational Definitions ───────────────────────────────────────────────
add_heading("5. Operational Definition of OEA Layers", level=1)
add_table(
    ["Layer", "Computational Meaning", "Mechanism", "Observable Effect"],
    [
        ["Ontological Anchoring", "Vocabulary domain constraint", "BM25 retrieval + vocab projection", "Higher in-distribution log-prob"],
        ["Epistemic Filtering", "Candidate quality discrimination", "K=3 candidates scored by log G_0; highest kept", "Improved TRR; reduced FRR"],
        ["Recursive Feedback", "Recursive persistence dynamics", "Output of step t conditions step t+1", "Stability retention over T iterations"],
        ["Calibration", "Directional shaping of distributions", "Selection by log G_0 score", "Correct direction improves TRR"],
        ["Miscalibration", "Anti-aligned shaping", "argmin replaces argmax", "TRR falls below control (H2)"],
        ["Recursive Drift", "Degradation across cycles", "Measured by 1 \u2212 S_t", "Higher drift = loss of fidelity"],
    ],
    "Table 2. Operational definitions of OEA layers and concepts."
)

# ── 6. Methodology ───────────────────────────────────────────────────────────
add_heading("6. Methodology", level=1)
add_body(
    "Each generation step produces output that conditions the next step, forming a feedback "
    "loop over a shared context window. The protocol has three layers: (1) Ontological "
    "Anchoring: BM25 retrieval + vocabulary projection, (2) Epistemic Filtering: K=3 "
    "candidate scoring under the frozen reference model, (3) Recursive Feedback: output of "
    "step t conditions step t+1 via context accumulation."
)
add_body(
    "E1 (Recursive Stability): 10 recursive generations per seed. Compare control vs OEA on "
    "JSD and stability. E2 (Epistemic Friction): Inject synthetic falsehoods; measure TRR "
    "and FRR. Execution: 30 seeded runs per condition."
)

# ── 7. Pilot Results ─────────────────────────────────────────────────────────
add_heading("7. Pilot Results", level=1)
add_table(
    ["Metric", "Control", "OEA"],
    [
        ["Stability score", "0.860 [0.841, 0.880]", "0.981 [0.978, 0.984]"],
        ["KL divergence", "2.660 [2.123, 3.197]", "0.287 [0.178, 0.397]"],
        ["TRR", "0.612 [0.603, 0.620]", "0.844 [0.838, 0.850]"],
        ["FRR", "0.233 [0.226, 0.239]", "0.121 [0.117, 0.126]"],
    ],
    "Table 3. Pilot experiment metrics (mean, 95% CI). Bigram proxy harness."
)

add_heading("7.1 Ablation Study", level=2)
add_body(
    "Miscalibration reversal: CQ < 0.5 flips rejection direction: FRR rises to 0.651, TRR "
    "falls to 0.257\u2014consistent with H2. ROUGE-L decline under anchoring: vocabulary "
    "anchoring concentrates token identity to the reference domain (improving log-probability) "
    "but does not preserve original phrase sequences. These metrics are orthogonal."
)
add_table(
    ["Variant", "Stability", "TRR", "FRR"],
    [
        ["control_replace", "0.449", "0.582", "0.241"],
        ["control_accumulate", "0.943", "0.582", "0.241"],
        ["ablation_E", "0.935", "0.752", "0.134"],
        ["ablation_EA", "0.950", "0.806", "0.100"],
        ["ablation_miscalibrated", "0.943", "0.257", "0.651"],
        ["oea_full", "0.937", "0.836", "0.081"],
    ],
    "Table 4. Ablation results (648 runs each, two domain corpora). d=4.56, p<0.001."
)

# ── 8. Real LLM Validation ──────────────────────────────────────────────────
add_heading("8. Real LLM Validation", level=1)
add_body(
    "Four models spanning three architecture families: distilgpt2 (82M), gpt2 (124M) from "
    "GPT-2, gpt-neo-125M (125M) from GPT-Neo (local attention), and Qwen2.5-1.5B (1.5B) "
    "from Qwen (RoPE, GQA, SwiGLU; September 2024). The Qwen model is 10\u00d7 larger with "
    "modern architecture, addressing the small-model artifact concern."
)
add_table(
    ["Model", "Variant", "Log-prob \u0394", "JSD", "ROUGE-L"],
    [
        ["distilgpt2", "control", "\u2014", "0.459", "0.046"],
        ["", "oea_anchored", "+0.62", "0.471", "0.038"],
        ["", "oea_miscalib.", "\u22121.14", "0.356", "0.054"],
        ["gpt2", "control", "\u2014", "0.363", "0.047"],
        ["", "oea_anchored", "+1.63", "0.438", "0.031"],
        ["", "oea_miscalib.", "\u22120.55", "0.384", "0.055"],
        ["gpt-neo", "control", "\u2014", "0.408", "0.049"],
        ["", "oea_anchored", "+0.82", "0.442", "0.037"],
        ["", "oea_miscalib.", "\u22121.37", "0.347", "0.054"],
        ["Qwen2.5", "control", "\u2014", "0.395", "0.054"],
        ["", "oea_anchored", "+0.88", "0.375", "0.068"],
        ["", "oea_miscalib.", "\u22120.62", "0.346", "0.050"],
    ],
    "Table 5. Real LLM validation (10 seeds \u00d7 10 iterations, CUDA 12.1, RTX 4070 SUPER)."
)
add_body(
    "H1 supported: oea_anchored improves log-probability by +0.62 to +1.63 nats across all "
    "four models. H2 supported: oea_miscalibrated degrades log-probability below control "
    "(\u22120.55 to \u22121.37 nats). H3 supported: oea_rag_only degrades log-probability "
    "(\u22120.07 to \u22120.54 nats). ROUGE-L dissociation is consistent with log-probability "
    "and ROUGE-L being orthogonal quality dimensions."
)

# ── 9. Recursive Memory Drift ────────────────────────────────────────────────
add_heading("9. Recursive Memory Drift Benchmark", level=1)
add_body(
    "A 30-step recursive memory drift benchmark compares oea_controlled vs uncontrolled over "
    "20 seeds in the bigram proxy harness. At step 30: OEA-controlled semantic drift JSD "
    "\u22120.006 lower (95% CI overlapping); entity retention \u0394 = \u22120.001 (negligible). "
    "The bigram harness provides only structural support for H4. Neural validation is required."
)

# ── 10. Baseline Competition ─────────────────────────────────────────────────
add_heading("10. Baseline Competition", level=1)
add_table(
    ["Method", "Mechanism", "Limitation"],
    [
        ["Temperature red.", "Global entropy suppression", "No epistemic discrimination"],
        ["Top-k", "Vocabulary truncation", "No directional calibration"],
        ["Entropy filter", "Uncertainty-based gating", "No semantic anchoring"],
        ["Repetition penalty", "Token damping", "No recursive coherence objective"],
        ["RAG-only", "Context injection", "No quality-aware filtering"],
        ["OEA", "Directional calibration + epistemic filter", "Designed for recursive coherence"],
    ],
    "Table 6. Mechanistic comparison: why OEA differs from generic decoding constraints."
)
add_body(
    "OEA achieves the highest TRR (0.746 vs next-best entropy_filter 0.630). Stability "
    "differences are not significant at N=20 (p > 0.3); real LLM baseline comparisons remain "
    "future work."
)

# ── 11. Discussion ───────────────────────────────────────────────────────────
add_heading("11. Discussion", level=1)
add_body(
    "Results are consistent with the primary thesis across four complementary experiments. "
    "The four-model, three-architecture validation supports the mechanism via log-probability "
    "improvement and ROUGE-L dissociation\u2014two orthogonal metrics. Cross-architecture "
    "consistency (GPT-2, GPT-Neo, Qwen2.5) and cross-scale consistency (82M to 1.5B) "
    "strengthen external validity."
)
add_body(
    "The key asymmetry: stability and epistemic reliability are partially orthogonal. "
    "Single-layer ablations can achieve comparable stability while exhibiting substantially "
    "weaker epistemic filtering."
)

# ── 12. Implications ─────────────────────────────────────────────────────────
add_heading("12. Implications for Recursive Systems", level=1)
add_bullet("Recursive planning drift: calibration-directed filtering may reduce plan degradation.")
add_bullet("Hallucinated refinement: miscalibrated correction signals amplify confident errors.")
add_bullet("Synthetic contamination in memory: ontological anchoring enforces domain boundaries.")
add_bullet("Specification degradation: recursive spec refinement is vulnerable to semantic drift.")
add_body(
    "These connections are suggestive, not demonstrated. We do not claim improvements to "
    "autonomous reasoning, planning competence, or general intelligence."
)

# ── 13. Failure Modes ────────────────────────────────────────────────────────
add_heading("13. Failure Modes and Limitations", level=1)
add_bullet("Bigram proxy scope: cannot model long-range dependencies or neural hallucination.")
add_bullet("Frozen-weights scope: not a substitute for fine-tuning experiments.")
add_bullet("Over-anchoring harms diversity: reduces ROUGE-L recall.")
add_bullet("CQ mismatch: measured CQ = 0.446 vs design estimate 0.83. ECE validation is future work.")
add_bullet("Scale: 82M to 1.5B parameters; extrapolation to >10B is not warranted.")
add_bullet("Baseline comparison power: N=20 seeds insufficient for significant pairwise comparisons.")
add_bullet("All causal language is provisional.")

# ── 14. Reproducibility ─────────────────────────────────────────────────────
add_heading("14. Reproducibility and Data Availability", level=1)
add_body(
    "All bigram experiments reproduce in <10 minutes (no GPU). Artifact integrity verified "
    "by SHA-256 hashes (experiments/manifest.json). Dockerfile provides containerized "
    "environment. All four real LLM models are validated and artifacts committed. "
    "Public-domain corpora: literary (Carroll, Austen, Melville, Hume, Darwin) and scientific "
    "(Newton, Feynman, Sagan). Code: https://github.com/BitConcepts/oea-framework-paper"
)

# ── 15. Conclusion ───────────────────────────────────────────────────────────
add_heading("15. Conclusion", level=1)
add_body(
    "We presented OEA as a testable, three-layer intervention hypothesis for improving "
    "reliability in recursive generative systems, with explicit non-claims: no AGI, no formal "
    "symbolic ontology, no autonomous agency, no causal proof."
)
add_body(
    "The full OEA protocol achieves TRR = 0.836, FRR = 0.081 (d = 4.56, p < 0.001) and "
    "improves mean log-probability by +0.62 to +1.63 nats across four models and three "
    "architecture families (82M\u20131.5B parameters). The anti-calibrated control provides "
    "strong ablation evidence for calibration direction as the operative variable."
)
add_body(
    "Future work: (1) frontier-scale model validation (>10B parameters); (2) recursive system "
    "integration with tool use and persistent memory; (3) ECE-based CQ validation; "
    "(4) neural recursive memory drift experiments."
)

# ── References ───────────────────────────────────────────────────────────────
add_heading("References", level=1)
refs = [
    "[1] Shumailov et al., \"AI models collapse when trained on recursively generated data,\" Nature, 2024.",
    "[2] Fu et al., \"Towards a theoretical understanding of synthetic data in LLM post-training,\" 2025.",
    "[3] Gerstgrasser et al., \"Is model collapse inevitable?\" arXiv:2402.07712, 2024.",
    "[4] Schaeffer et al., \"Position: LLM training on model collapse,\" ICML, 2025.",
    "[5] Popper, The Logic of Scientific Discovery, 1959.",
    "[6] Drayson et al., \"Machine-generated text detection prevents collapse,\" EMNLP, 2025.",
    "[7] Zhu et al., \"Synthesize, don't collapse,\" ICML, 2025.",
    "[8] Kov\u00e1\u010d et al., \"Recursive distributional shift under lexical diversity,\" 2025.",
    "[9] Keisha et al., \"Knowledge collapse in LLMs,\" NeurIPS Workshop, 2025.",
    "[10] Lewis et al., \"Retrieval-augmented generation for knowledge-intensive NLP,\" NeurIPS, 2020.",
    "[11] Abbasi Yadkori et al., \"To believe or not to believe your LLM,\" Google DeepMind, 2024.",
    "[12] Guo et al., \"On calibration of modern neural networks,\" ICML, 2017.",
    "[13] Roumeliotis et al., \"Trust-aware LLM orchestration,\" arXiv:2507.10571, 2025.",
    "[14] Fu et al., \"Self-verification improves few-shot clinical NLP,\" NeurIPS, 2025.",
]
for ref in refs:
    p = doc.add_paragraph(ref)
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.left_indent = Pt(18)
    p.paragraph_format.space_after = Pt(2)
    for run in p.runs:
        run.font.size = Pt(9)

# ── Save ─────────────────────────────────────────────────────────────────────
doc.save(str(OUT))
print(f"Done. Saved to: {OUT}")
print(f"Size: {OUT.stat().st_size // 1024} KB")
